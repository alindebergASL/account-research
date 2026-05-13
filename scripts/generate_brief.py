#!/usr/bin/env python3
"""Render a structured account brief (JSON) into PDF and DOCX.

Usage:
    python3 scripts/generate_brief.py <brief.json> [--out-dir out]

Prints the absolute paths of the generated files (one per line) to stdout.
"""
from __future__ import annotations

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
)


SECTIONS = [
    ("Account Snapshot", "snapshot"),
    ("Segment / Industry", "segment"),
    ("Account Priority Summary", "priority_summary"),
    ("Recent Strategic Signals", "recent_signals"),
    ("AI / Technology Maturity Rating", "ai_tech_maturity"),
    ("Top Business or Technology Initiatives", "top_initiatives"),
    ("Technical Footprint", "technical_footprint"),
    ("Programs & Procurement Signals", "programs_procurement"),
    ("Key Personas and Conversation Openers", "personas"),
    ("Buying / Procurement / Decision Path", "buying_path"),
    ("Recommended First Conversation Angle", "first_angle"),
    ("Risks and Watch-Outs", "risks"),
    ("Competitive / Partner / Vendor Signals", "competitive_signals"),
    ("Recommended Next Action", "next_action"),
    ("Insights", "extensions"),
    ("Key Sources", "sources"),
]

TECH_ROWS = [
    ("AI / automation in production", "ai_in_production", "list"),
    ("Active pilots / POCs", "active_pilots", "list"),
    ("Cloud platforms", "cloud_platforms", "list"),
    ("Data infrastructure", "data_infrastructure", "text"),
    ("Clinical platforms / EHR", "clinical_platforms", "text"),
    ("Analytics / BI stack", "analytics_bi_stack", "text"),
    ("Build vs. buy posture", "build_vs_buy_posture", "text"),
    ("Competitive incumbents", "competitive_incumbents", "list"),
]

PROGRAM_ROWS = [
    ("Modernization grants", "modernization_grants", "list"),
    ("Consortium / cooperative purchasing", "consortium_purchasing", "list"),
    ("Active RFPs / contracts (12-18 mo)", "active_rfps_contracts", "list"),
    ("AI governance / responsible AI policy", "ai_governance_policy", "text"),
    ("Publicly stated AI use cases", "public_ai_use_cases", "list"),
]


def _format_value(value, kind):
    if kind == "list":
        items = [clean(x) for x in (value or []) if clean(x)]
        return "; ".join(items) if items else "Not found in public sources."
    text = clean(value)
    return text if text else "Not found in public sources."


_MD_ARTIFACTS = re.compile(r"(\*\*|`{1,3}|^#{1,6}\s+)", re.MULTILINE)


def clean(text):
    if text is None:
        return ""
    if not isinstance(text, str):
        text = str(text)
    return _MD_ARTIFACTS.sub("", text).strip()


def slugify(name: str) -> str:
    s = re.sub(r"[^A-Za-z0-9]+", "-", name).strip("-").lower()
    return s or "account"


def maturity_line(m):
    if not isinstance(m, dict):
        return clean(m)
    rating = m.get("rating", "?")
    rationale = clean(m.get("rationale", ""))
    return f"Rating {rating} / 5 — {rationale}" if rationale else f"Rating {rating} / 5"



def extension_label(ext):
    return clean(ext.get("kind", "insight")).title()


# List items accept either a plain string (legacy / pre-PR-A) or a
# {heading?, text} object (PR-A spec). Renderers fold both into one
# display string of the form "Heading: text" when heading is present.
def list_item_text(item):
    if isinstance(item, dict):
        heading = clean(item.get("heading"))
        text = clean(item.get("text"))
        if heading and text:
            return f"{heading}: {text}"
        return heading or text
    return clean(item)


def add_docx_extensions(doc, extensions):
    for ext in extensions:
        if not isinstance(ext, dict):
            continue
        title = clean(ext.get("title")) or "Insight"
        meta = " · ".join([x for x in [extension_label(ext), clean(ext.get("confidence")), "Added in chat" if ext.get("source") == "chat" else ""] if x])
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        if meta:
            mr = p.add_run(f" ({meta})")
            mr.font.size = Pt(9)
            mr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        kind = ext.get("kind")
        if kind == "table":
            columns = [clean(c) for c in ext.get("columns", [])]
            rows = ext.get("rows", []) if isinstance(ext.get("rows"), list) else []
            if columns:
                tbl = doc.add_table(rows=1, cols=len(columns))
                tbl.style = "Light Grid Accent 1"
                for i, col in enumerate(columns):
                    tbl.rows[0].cells[i].text = col
                for row_values in rows:
                    row = tbl.add_row().cells
                    for i in range(len(columns)):
                        row[i].text = clean(row_values[i] if isinstance(row_values, list) and i < len(row_values) else "")
        elif kind == "list":
            for item in ext.get("items", []) or []:
                doc.add_paragraph(list_item_text(item), style="List Bullet")
        elif kind == "card":
            para = doc.add_paragraph()
            run = para.add_run(clean(ext.get("body")))
            run.bold = False
            badges = ext.get("badges") if isinstance(ext.get("badges"), list) else []
            if badges:
                badges_p = doc.add_paragraph(" · ".join(clean(b) for b in badges if clean(b)))
                if badges_p.runs:
                    badges_p.runs[0].font.size = Pt(9)
                    badges_p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        else:
            doc.add_paragraph(clean(ext.get("body")))

        if clean(ext.get("why_included")):
            why = doc.add_paragraph(f"Why included: {clean(ext.get('why_included'))}")
            why.runs[0].font.size = Pt(9)
            why.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def pdf_extension_flowables(ext, styles, page_w):
    flow = []
    if not isinstance(ext, dict):
        return flow
    title = clean(ext.get("title")) or "Insight"
    bits = [extension_label(ext), clean(ext.get("confidence"))]
    if ext.get("source") == "chat":
        bits.append("Added in chat")
    flow.append(Paragraph(f"<b>{title}</b> — {' · '.join([b for b in bits if b])}", styles["body"]))
    kind = ext.get("kind")
    if kind == "table":
        cols = [clean(c) for c in ext.get("columns", [])]
        rows = ext.get("rows", []) if isinstance(ext.get("rows"), list) else []
        if cols:
            data = [[cell for cell in cols]]
            for row in rows:
                data.append([clean(row[i] if isinstance(row, list) and i < len(row) else "") for i in range(len(cols))])
            col_w = [page_w / len(cols)] * len(cols)
            flow.append(_table([[Paragraph(clean(c), styles["tbl"]) for c in r] for r in data], col_w))
    elif kind == "list":
        for item in ext.get("items", []) or []:
            flow.append(Paragraph(f"• {list_item_text(item)}", styles["bullet"]))
    elif kind == "card":
        body = Paragraph(clean(ext.get("body")), styles["body"])
        cell_flows = [body]
        badges = ext.get("badges") if isinstance(ext.get("badges"), list) else []
        if badges:
            cell_flows.append(
                Paragraph(" · ".join(clean(b) for b in badges if clean(b)), styles["meta"])
            )
        t = Table([[cell_flows]], colWidths=[page_w])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f5f7fb")),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ]))
        flow.append(t)
    else:
        flow.append(Paragraph(clean(ext.get("body")), styles["body"]))
    if clean(ext.get("why_included")):
        flow.append(Paragraph(f"Why included: {clean(ext.get('why_included'))}", styles["meta"]))
    return flow

# ---------- DOCX ----------

def render_docx(brief: dict, path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(f"Account Brief: {clean(brief.get('account_name', 'Unknown'))}")
    run.bold = True
    run.font.size = Pt(20)

    meta_bits = []
    if brief.get("segment"):
        meta_bits.append(clean(brief["segment"]))
    if brief.get("audience"):
        meta_bits.append(f"Audience: {clean(brief['audience'])}")
    meta_bits.append(f"Generated: {clean(brief.get('generated_at', datetime.utcnow().strftime('%Y-%m-%d')))}")
    meta = doc.add_paragraph(" · ".join(meta_bits))
    meta.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    meta.runs[0].font.size = Pt(9)

    for heading, key in SECTIONS:
        value = brief.get(key)
        if value in (None, "", [], {}):
            continue
        doc.add_heading(heading, level=1)

        if key == "recent_signals" and isinstance(value, list):
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Light Grid Accent 1"
            hdr = tbl.rows[0].cells
            hdr[0].text = "Signal"
            hdr[1].text = "Source"
            hdr[2].text = "Confidence"
            for item in value:
                row = tbl.add_row().cells
                row[0].text = clean(item.get("text"))
                row[1].text = clean(item.get("source"))
                row[2].text = clean(item.get("confidence"))

        elif key == "ai_tech_maturity":
            doc.add_paragraph(maturity_line(value))

        elif key == "top_initiatives" and isinstance(value, list):
            tbl = doc.add_table(rows=1, cols=4)
            tbl.style = "Light Grid Accent 1"
            hdr = tbl.rows[0].cells
            hdr[0].text = "Initiative"
            hdr[1].text = "Detail"
            hdr[2].text = "Confidence"
            hdr[3].text = "Source"
            for item in value:
                row = tbl.add_row().cells
                row[0].text = clean(item.get("title"))
                row[1].text = clean(item.get("detail"))
                row[2].text = clean(item.get("confidence"))
                row[3].text = clean(item.get("source"))

        elif key == "technical_footprint" and isinstance(value, dict):
            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = "Light Grid Accent 1"
            hdr = tbl.rows[0].cells
            hdr[0].text = "Dimension"
            hdr[1].text = "Detail"
            for label, fld, kind in TECH_ROWS:
                if fld == "clinical_platforms" and not clean(value.get(fld)):
                    continue
                row = tbl.add_row().cells
                row[0].text = label
                row[1].text = _format_value(value.get(fld), kind)

        elif key == "programs_procurement" and isinstance(value, dict):
            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = "Light Grid Accent 1"
            hdr = tbl.rows[0].cells
            hdr[0].text = "Program area"
            hdr[1].text = "Detail"
            for label, fld, kind in PROGRAM_ROWS:
                row = tbl.add_row().cells
                row[0].text = label
                row[1].text = _format_value(value.get(fld), kind)

        elif key == "personas" and isinstance(value, list):
            tbl = doc.add_table(rows=1, cols=6)
            tbl.style = "Light Grid Accent 1"
            hdr = tbl.rows[0].cells
            for i, h in enumerate(["Name", "Title", "Priority", "Opener", "Confidence", "Source"]):
                hdr[i].text = h
            for p in value:
                row = tbl.add_row().cells
                row[0].text = clean(p.get("name")) or "—"
                row[1].text = clean(p.get("title"))
                row[2].text = clean(p.get("priority"))
                row[3].text = clean(p.get("opener"))
                row[4].text = clean(p.get("confidence"))
                row[5].text = clean(p.get("source"))

        elif key == "extensions" and isinstance(value, list):
            add_docx_extensions(doc, value)

        elif key == "sources" and isinstance(value, list):
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Light Grid Accent 1"
            hdr = tbl.rows[0].cells
            hdr[0].text = "Title"
            hdr[1].text = "URL"
            hdr[2].text = "Accessed"
            for s in value:
                row = tbl.add_row().cells
                row[0].text = clean(s.get("title"))
                row[1].text = clean(s.get("url"))
                row[2].text = clean(s.get("accessed"))

        elif key in ("risks", "competitive_signals") and isinstance(value, list):
            for item in value:
                doc.add_paragraph(clean(item), style="List Bullet")

        else:
            doc.add_paragraph(clean(value))

    doc.save(str(path))


# ---------- PDF ----------

def _pdf_styles():
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle("title", parent=base["Title"], fontSize=20, leading=24, spaceAfter=4),
        "meta": ParagraphStyle("meta", parent=base["Normal"], fontSize=9, textColor=colors.grey, spaceAfter=14),
        "h1": ParagraphStyle("h1", parent=base["Heading1"], fontSize=13, leading=16, spaceBefore=12, spaceAfter=6, textColor=colors.HexColor("#1f3a5f")),
        "body": ParagraphStyle("body", parent=base["Normal"], fontSize=10.5, leading=14, spaceAfter=6),
        "bullet": ParagraphStyle("bullet", parent=base["Normal"], fontSize=10.5, leading=14, leftIndent=14, bulletIndent=2, spaceAfter=2),
        "tbl": ParagraphStyle("tbl", parent=base["Normal"], fontSize=9, leading=11),
    }
    return styles


def _table(data, col_widths):
    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f3a5f")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cccccc")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f7fb")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


def render_pdf(brief: dict, path: Path) -> None:
    styles = _pdf_styles()
    doc = SimpleDocTemplate(
        str(path), pagesize=LETTER,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
        title=f"Account Brief — {brief.get('account_name', '')}",
    )
    story = []

    story.append(Paragraph(clean(f"Account Brief: {brief.get('account_name', 'Unknown')}"), styles["title"]))
    meta_bits = []
    if brief.get("segment"):
        meta_bits.append(clean(brief["segment"]))
    if brief.get("audience"):
        meta_bits.append(f"Audience: {clean(brief['audience'])}")
    meta_bits.append(f"Generated: {clean(brief.get('generated_at', datetime.utcnow().strftime('%Y-%m-%d')))}")
    story.append(Paragraph(" &middot; ".join(meta_bits), styles["meta"]))

    def cell(text):
        return Paragraph(clean(text).replace("\n", "<br/>"), styles["tbl"])

    page_w = LETTER[0] - 1.4 * inch

    for heading, key in SECTIONS:
        value = brief.get(key)
        if value in (None, "", [], {}):
            continue
        story.append(Paragraph(heading, styles["h1"]))

        if key == "recent_signals" and isinstance(value, list):
            data = [["Signal", "Source", "Confidence"]]
            for it in value:
                data.append([cell(it.get("text")), cell(it.get("source")), cell(it.get("confidence"))])
            story.append(_table(data, [page_w * 0.55, page_w * 0.30, page_w * 0.15]))

        elif key == "ai_tech_maturity":
            story.append(Paragraph(clean(maturity_line(value)), styles["body"]))

        elif key == "top_initiatives" and isinstance(value, list):
            data = [["Initiative", "Detail", "Confidence", "Source"]]
            for it in value:
                data.append([cell(it.get("title")), cell(it.get("detail")), cell(it.get("confidence")), cell(it.get("source"))])
            story.append(_table(data, [page_w * 0.22, page_w * 0.45, page_w * 0.13, page_w * 0.20]))

        elif key == "technical_footprint" and isinstance(value, dict):
            data = [["Dimension", "Detail"]]
            for label, fld, kind in TECH_ROWS:
                if fld == "clinical_platforms" and not clean(value.get(fld)):
                    continue
                data.append([cell(label), cell(_format_value(value.get(fld), kind))])
            story.append(_table(data, [page_w * 0.32, page_w * 0.68]))

        elif key == "programs_procurement" and isinstance(value, dict):
            data = [["Program area", "Detail"]]
            for label, fld, kind in PROGRAM_ROWS:
                data.append([cell(label), cell(_format_value(value.get(fld), kind))])
            story.append(_table(data, [page_w * 0.32, page_w * 0.68]))

        elif key == "personas" and isinstance(value, list):
            data = [["Name", "Title", "Priority", "Opener", "Conf.", "Source"]]
            for p in value:
                data.append([
                    cell(p.get("name") or "—"),
                    cell(p.get("title")),
                    cell(p.get("priority")),
                    cell(p.get("opener")),
                    cell(p.get("confidence")),
                    cell(p.get("source")),
                ])
            story.append(_table(data, [page_w * 0.14, page_w * 0.16, page_w * 0.14, page_w * 0.32, page_w * 0.10, page_w * 0.14]))

        elif key == "extensions" and isinstance(value, list):
            for ext in value:
                story.extend(pdf_extension_flowables(ext, styles, page_w))

        elif key == "sources" and isinstance(value, list):
            data = [["Title", "URL", "Accessed"]]
            for s in value:
                data.append([cell(s.get("title")), cell(s.get("url")), cell(s.get("accessed"))])
            story.append(_table(data, [page_w * 0.35, page_w * 0.50, page_w * 0.15]))

        elif key in ("risks", "competitive_signals") and isinstance(value, list):
            for item in value:
                story.append(Paragraph(f"• {clean(item)}", styles["bullet"]))

        else:
            story.append(Paragraph(clean(value), styles["body"]))

        story.append(Spacer(1, 4))

    doc.build(story)


def main(argv=None) -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("brief_json", help="Path to brief JSON file")
    ap.add_argument("--out-dir", default="out", help="Output directory")
    ap.add_argument("--formats", default="pdf,docx", help="Comma-separated: pdf,docx")
    args = ap.parse_args(argv)

    src = Path(args.brief_json).resolve()
    if not src.exists():
        print(f"error: {src} not found", file=sys.stderr)
        return 2
    brief = json.loads(src.read_text(encoding="utf-8"))

    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    name = brief.get("account_name") or "account"
    stamp = datetime.utcnow().strftime("%Y%m%d")
    slug = slugify(name)
    base = out_dir / f"{slug}-{stamp}"

    formats = {f.strip().lower() for f in args.formats.split(",") if f.strip()}
    written = []

    if "docx" in formats:
        docx_path = base.with_suffix(".docx")
        render_docx(brief, docx_path)
        written.append(docx_path)

    if "pdf" in formats:
        pdf_path = base.with_suffix(".pdf")
        render_pdf(brief, pdf_path)
        written.append(pdf_path)

    for p in written:
        print(str(p))
    return 0


if __name__ == "__main__":
    sys.exit(main())
